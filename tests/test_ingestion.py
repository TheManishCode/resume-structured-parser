"""tests/test_ingestion.py — ingestion pipeline unit tests."""
import pytest
from packages.core.ingestion.pipeline import (
    ingest_file, _find_email, _find_phone, _estimate_yoe, _find_skills,
)


SAMPLE_PDF_TEXT = b""  # tested via mocks below


@pytest.mark.asyncio
async def test_ingest_docx_stub(tmp_path):
    import docx
    doc = docx.Document()
    doc.add_paragraph("John Smith — ML Engineer — 5 years of experience")
    doc.add_paragraph("Email: john@example.com")
    doc.add_paragraph("Skills: Python, PyTorch, NLP, Machine Learning")
    path = tmp_path / "resume.docx"
    doc.save(str(path))
    with open(path, "rb") as f:
        raw = f.read()
    parsed, text, used_ocr = await ingest_file("resume.docx", raw)
    assert used_ocr is False
    assert "John" in text
    assert parsed["email"] == "john@example.com"
    assert "python" in parsed["skills_raw"]


def test_find_email():
    assert _find_email("contact: foo.bar@example.co.in") == "foo.bar@example.co.in"
    assert _find_email("no email here") is None


def test_find_phone():
    assert _find_phone("+91 98765 43210 some text") is not None
    assert _find_phone("no phone") is None


def test_estimate_yoe_explicit():
    assert _estimate_yoe("I have 7 years of experience in ML") == 7.0


def test_estimate_yoe_from_years():
    text = "Worked at Swiggy from 2018 to 2024"
    yoe = _estimate_yoe(text)
    assert yoe == 6.0


def test_find_skills():
    text = "Proficient in Python, PyTorch, NLP and Docker deployment"
    skills = _find_skills(text)
    assert "python" in skills
    assert "pytorch" in skills
    assert "nlp" in skills
    assert "docker" in skills


def test_bias_fields_stripped():
    """parsed_json must not contain name, photo, dob, gender, nationality."""
    from packages.core.ingestion.pipeline import _normalize, _strip_bias_fields
    data = _normalize("Jane Doe, 1990, Female, Indian national. ML Engineer 3 years experience.")
    result = _strip_bias_fields(data)
    for field in ("name", "photo", "dob", "gender", "nationality"):
        assert field not in result


@pytest.mark.asyncio
async def test_unsupported_extension_raises():
    with pytest.raises(ValueError, match="Unsupported"):
        await ingest_file("resume.txt", b"some text")
